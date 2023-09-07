
from lib.SQL import CONN, CURSOR
from datetime import datetime
from lib.overview import Overview
from lib.GPTcontainer import get_completion
from bs4 import BeautifulSoup
import ast
from docx import Document


class Conversation:
    def __init__(self):
        self.overview = Overview()
        self.conversation_id = self.overview.initialize_convo()
        Conversation.create_table()

    # create reference in overview table with id

    
    @classmethod
    def get_highest_id(cls):
        CURSOR.execute("SELECT MAX(convo_id) FROM conversations")
        highest_id = CURSOR.fetchone()[0]
        return highest_id if highest_id else 0

  
    @classmethod
    def create_table(cls):
        query = """
        CREATE TABLE IF NOT EXISTS conversations (
        id INTEGER PRIMARY KEY,
        time TIMESTAMP,
        convo_id INTEGER,
        content TEXT
        )
        """
        CURSOR.execute(query)


    def add_row(self, content):
        current_time = datetime.now().strftime('%H:%M:%S')
        query = """
        INSERT INTO conversations
        (time, convo_id, content) VALUES (?,?,?)
        """
        CURSOR.execute(query, [current_time, self.conversation_id, content])
        CONN.commit()

    @classmethod
    def reset_table(cls, table):
        query = f"DROP TABLE {table}"
        CURSOR.execute(query)
        CONN.commit()

    def end_conversation(self):
        print("ending convo")
        convo = Overview.get_readable_conversation(self.overview.id)

        print(f"conversation: {convo}")

        prompt = """
        Provide a title and summary for the following transcript formatted as a Python dictionary with no other text. 
        Keep in mind that the transcript may be imperfect.
        Your output should look like this with your text:
        {
        'title': 'Placeholder Title',
        'summary': 'This is a placeholder summary.'
        } \n
        """ + convo

        response = get_completion(prompt)


        try:
            titleSummary = ast.literal_eval(response)
        except:
            titleSummary = None
        
        self.overview.add_title_summary(titleSummary)


    @classmethod
    def reset_database(cls):
        cls.reset_table("conversations")
        cls.reset_table("convo_overview")

    @classmethod
    def delete_convo(cls, id):
        query1 = "DELETE FROM conversations WHERE convo_id = ?"
        query2 = "DELETE FROM convo_overview WHERE id = ?"
        CURSOR.execute(query1, [id])
        CURSOR.execute(query2, [id])
        CONN.commit()

    @classmethod
    def export(cls, id, raw_html=False):
        print("Exporting...")
        convo = Overview.get_readable_conversation(id)


        prompt = """
        Given the following audio transcript, create an HTML document summarizing the transcript. Only respond with the HTML and no other text. Use the following formatting:
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Summary</title>
</head>
<body>

<h1>[create an appropriate title for the transcript and insert it here]</h1>

<p>[insert a one sentence summary of the transcript here]</p>

<h2>Main Topics Discussed</h2>
<ul>
    <li>[Insert topic 1 here]</li>
    <li>[Insert topic 2 here]</li>
    <!-- Add as many topics as needed. -->
</ul>

<h2>Key Takeaways</h2>
<ul>
    <li>[Insert takeaway 1 here]</li>
    <li>[Insert takeaway 2 here]</li>
    <!-- Add as many takeaways as needed. -->
</ul>

<!-- Notes for detailed summary: -->
<!-- Include all essential information, such as vocabulary terms and key concepts -->
<!-- Strictly base your notes on the provided information, without adding any external information. -->
<h2>Detailed Summary</h2>
<p>[create detailed notes about the transcript]</p>

<!-- Optional: Only include the section below if there are action items -->
<h2>Next Steps/Action Items</h2>
<ul>
    <li>[Action 1]</li>
    <li>[Action 2]</li>
    <!-- Add as many actions as needed. -->
</ul>

</body>
</html>

TRANSCRIPT (may contain errors):
        """ + convo
        html_text = get_completion(prompt)

        if raw_html:
            return html_text
       
        document = Document()
        soup = BeautifulSoup(html_text, 'html.parser')
        for element in soup.body:
            if element.name and element.name.startswith('h') and element.name[1:].isdigit():
                level = int(element.name[1:]) - 1
                document.add_heading(element.text, level=level)

            elif element.name == 'p':
                document.add_paragraph(element.text)
            elif element.name == 'ul':
                for item in element.find_all('li'):
                    document.add_paragraph(item.text, style='ListBullet')
            elif element.name == 'ol':
                for item in element.find_all('li'):
                    document.add_paragraph(item.text, style='ListNumber')

        document.save(f"exports/Overview_for_convo_{id}.docx")
        return None






        









        
